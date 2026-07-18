"""PDF and DOCX readers: one Document per PDF page, one per DOCX file."""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path

import docx
import pdfplumber

from ..contracts import Document, Dossier, SourceRef

logging.getLogger("pdfminer").setLevel(logging.ERROR)

_UMLAUTS = str.maketrans({"\xc4": "AE", "\xd6": "OE", "\xdc": "UE"})

# Priority-ordered bilingual DE+EN kind patterns, matched against the
# umlaut-folded uppercase filename plus leading document text.
_KINDS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("financial_statements",
     re.compile(r"BILANZ|(?<![A-Z])GUV(?![A-Z])|BALANCE\s*SHEET|INCOME\s*STATEMENT"
                r"|P\s*&\s*L|JAHRESABSCHLUSS|ANNUAL\s*REPORT")),
    ("audit_plan", re.compile(r"PRUEFUNGSPLAN|AUDIT\s*PLAN|(?<![A-Z])JET(?![A-Z])")),
    ("export_protocol", re.compile(r"EXPORTPROTOKOLL|EXPORT\s*LOG|GDPDU")),
    ("it_confirmation", re.compile(r"VOLLSTAENDIGKEIT|IT-?BESTAETIGUNG|COMPLETENESS")),
    ("bank_confirmation",
     re.compile(r"BANKBESTAETIGUNG|BANK\s*CONFIRMATION|SALDENBESTAETIGUNG")),
    ("bank_statement", re.compile(r"KONTOAUSZUG|BANK\s*STATEMENT|SETTLEMENT")),
)


def _classify(filename: str, text: str) -> str:
    hay = (filename + "\n" + text[:4000]).upper().translate(_UMLAUTS)
    for kind, pattern in _KINDS:
        if pattern.search(hay):
            return kind
    return "document"


def _load_pdf(path: Path, rel: str, dossier: Dossier) -> None:
    pages: list[tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                pages.append((page_no, text))
    if not pages:
        dossier.unparsed.append((rel, "no extractable text (image-only or empty)"))
        return
    kind = _classify(path.name, "\n".join(text for _, text in pages))
    for page_no, text in pages:
        dossier.documents.append(Document(
            kind=kind,
            ref=f"{path.stem}#p{page_no}",
            source=SourceRef(file=rel, page=page_no, excerpt=text[:200]),
            fields={"text": text},
        ))


def _load_docx(path: Path, rel: str, dossier: Dossier) -> None:
    parsed = docx.Document(str(path))
    lines = [para.text for para in parsed.paragraphs]
    for table in parsed.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    text = "\n".join(lines)
    if not text.strip():
        dossier.unparsed.append((rel, "no extractable text"))
        return
    dossier.documents.append(Document(
        kind=_classify(path.name, text),
        ref=path.stem,
        source=SourceRef(file=rel, page=1, excerpt=text[:200]),
        fields={"text": text},
    ))


def load_docs(root: Path, dossier: Dossier) -> None:
    """Scan root recursively for *.pdf / *.docx and append Document objects.

    Never raises: unreadable files land in dossier.unparsed instead.
    """
    root = Path(root)
    found: list[Path] = []
    for dirpath, _dirnames, filenames in os.walk(root, followlinks=True):
        for name in filenames:
            if name.startswith(("~$", ".")):
                continue
            if name.lower().endswith((".pdf", ".docx")):
                found.append(Path(dirpath) / name)
    for path in sorted(found):
        try:
            rel = path.relative_to(root).as_posix()
        except ValueError:
            rel = path.name
        try:
            if path.suffix.lower() == ".pdf":
                _load_pdf(path, rel, dossier)
            else:
                _load_docx(path, rel, dossier)
        except Exception as exc:
            dossier.unparsed.append((rel, f"{type(exc).__name__}: {exc}"))
